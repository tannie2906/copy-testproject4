import base64
from datetime import datetime, timezone
import mimetypes
from pathlib import Path
import random
import re
import shutil
from rest_framework.views import APIView
from rest_framework.response import Response #handle error
from rest_framework.permissions import IsAuthenticated, AllowAny
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.authtoken.models import Token
from rest_framework import viewsets, permissions, status
from rest_framework.decorators import api_view, permission_classes
from rest_framework import status
from rest_framework.generics import ListAPIView
from rest_framework.exceptions import ValidationError
from rest_framework.generics import DestroyAPIView
from rest_framework import generics, filters
from rest_framework.pagination import PageNumberPagination

from django.db.models.signals import post_save
from django.dispatch import receiver
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from django.views.decorators.http import require_GET
from django.core.files.storage import default_storage
from django.core.mail import send_mail

from django.views import View
from django.contrib.auth import authenticate
from django.contrib.auth.models import User
from django.http import JsonResponse, FileResponse
from django.http import HttpResponse, Http404
from django.conf import settings
from django.shortcuts import get_object_or_404
from django.urls import reverse
from django.core.exceptions import ObjectDoesNotExist
from django.utils.timezone import now 
from django.utils import timezone
from django.utils.text import slugify
from django.contrib.auth.password_validation import validate_password
from django.utils.decorators import method_decorator
from django.db import transaction

#from myapp.gmail_api import send_email
from testproject.settings import EMAIL_HOST_USER


from .models import DeletedFile, UploadedFile, File, SharedFile, Profile, Folder
from .serializers import DeletedFilesSerializer, UserSerializer, UploadedFileSerializer, UserRegistrationSerializer, FileSerializer, ProfilePictureSerializer, ProfileSerializer, FolderSerializer
from myapp.models import File, DeletedFile
from django.contrib.auth import get_user_model
from django_otp.plugins.otp_totp.models import TOTPDevice
from django.contrib.auth.decorators import login_required
from django.views.decorators.csrf import ensure_csrf_cookie
from rest_framework.authentication import TokenAuthentication
from django.middleware.csrf import get_token
from secrets import token_hex
import traceback
import qrcode
from io import BytesIO
from datetime import timedelta
from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django_otp.decorators import otp_required
from django_otp.plugins.otp_totp.models import TOTPDevice
from django.contrib.auth import login, authenticate
import pyqrcode
from django.views.decorators.http import require_POST
import base64
import pyotp
from rest_framework.parsers import JSONParser
from django.utils.crypto import get_random_string
from django.contrib.auth.hashers import make_password
from email.mime.text import MIMEText
from googleapiclient.discovery import build
from get_gmail_credentials import get_credentials 
from .serializers import PasswordResetRequestSerializer, PasswordResetConfirmSerializer
from django.contrib.auth.tokens import default_token_generator
from django.utils.http import urlsafe_base64_encode, urlsafe_base64_decode
from django.core.mail import send_mail
from myapp.gmail_api import send_reset_email 
from django.utils.http import urlsafe_base64_encode
from django.utils.encoding import force_bytes
from myapp.models import AuditLog
from django.core.validators import validate_email
from django.core.exceptions import ValidationError 
from .models import Lockbox

from .gmail_api import send_email_via_gmail  # This imports the function from gmail_api.py

import json
import logging
import uuid
import os
from urllib.parse import urljoin
from myapp.encryption_utils import decrypt_and_serve_file, encrypt_and_save_file, decrypt_file_to_temp
from django.core.files.uploadedfile import InMemoryUploadedFile, TemporaryUploadedFile
from cryptography.fernet import Fernet
from docx import Document
from rest_framework.response import Response
import uuid
from datetime import timedelta
from django.utils.timezone import now

logger = logging.getLogger(__name__)
User = get_user_model()

# Helper to validate file ownership
def validate_file_owner(file_id, user):
    try:
        file = File.objects.get(id=file_id)
        if file.user != user:
            raise ValidationError("You are not authorized to access this file.")
        return file
    except File.DoesNotExist:
        raise ValidationError("File not found.")

# ViewSet for files
class FileUploadView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        try:
            uploaded_files = request.FILES.getlist('files')  # Only handle file uploads
            user_id = request.user.id

            # Allowed file extensions and max size
            allowed_extensions = {'jpg', 'jpeg', 'png', 'pdf', 'txt', 'docx', 'xlsx', 'csv', 'zip'}
            max_size = 20 * 1024 * 1024  # 20 MB

            uploaded_file_data = []

            for uploaded_file in uploaded_files:
                file_extension = uploaded_file.name.split('.')[-1].lower()
                if file_extension not in allowed_extensions:
                    return Response({"error": f"Invalid file type: {uploaded_file.name}"}, status=400)
                if uploaded_file.size > max_size:
                    return Response({"error": f"File size exceeds 20MB: {uploaded_file.name}"}, status=400)

                # Save the uploaded file
                base_upload_dir = os.path.join(settings.MEDIA_ROOT, 'uploads')
                os.makedirs(base_upload_dir, exist_ok=True)

                safe_file_name = f"{slugify(uploaded_file.name)}.{file_extension}"
                file_path = os.path.join(base_upload_dir, safe_file_name)

                # Handle encryption based on file type
                if isinstance(uploaded_file, InMemoryUploadedFile):
                    uploaded_file.seek(0)  # Ensure reading from the start
                    encrypt_and_save_file(uploaded_file, file_path)

                elif isinstance(uploaded_file, TemporaryUploadedFile):
                    with open(uploaded_file.temporary_file_path(), 'rb') as temp_file:
                        encrypt_and_save_file(temp_file, file_path)

                # Save file metadata (No folder association)
                file_instance = File.objects.create(
                    file=f"uploads/{safe_file_name}",
                    file_name=uploaded_file.name,
                    size=uploaded_file.size,
                    user_id=user_id
                )

                uploaded_file_data.append({
                    "file_id": file_instance.id,
                    "file_name": file_instance.file_name
                })

            return Response({
                "message": "Files uploaded successfully!",
                "files": uploaded_file_data
            }, status=201)

        except Exception as e:
            print("Upload Error:", str(e))
            return Response({"error": str(e)}, status=500)

 
 # ViewSet for files
class FolderListView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, *args, **kwargs):
        # Fetch all folders
        folders = Folder.objects.filter(user=request.user)
        
        # Fetch only unlocked files
        files = File.objects.filter(user_id=request.user.id, is_deleted=False, is_locked=False)

        # Serialize the data
        folder_serializer = FolderSerializer(folders, many=True)
        file_serializer = FileSerializer(files, many=True)

        return Response({
            'folders': folder_serializer.data,
            'files': file_serializer.data  # Now includes only unlocked files
        })

# User Authentication and login
class CustomAuthToken(APIView):
    permission_classes = [AllowAny]

    def post(self, request):
        user = authenticate(
            username=request.data.get('username'),
            password=request.data.get('password')
        )
        if user:
            token, _ = Token.objects.get_or_create(user=user)
            return Response({"token": token.key})
        return Response({"error": "Invalid credentials"}, status=400)

#password reset 
@api_view(['POST'])
@permission_classes([])  # You can specify permissions if needed
def password_reset_request(request):
    """Send a password reset email via Gmail API."""
    email = request.data.get('email')  # Use request.data for JSON requests
    try:
        user = User.objects.get(email=email)  # Find the user by email
        token = default_token_generator.make_token(user)  # Generate a reset token
        uid = urlsafe_base64_encode(force_bytes(user.pk))  # Properly encode user ID
        reset_url = f'https://localhost:4200/reset-password/{uid}/{token}/'  # Construct the reset URL

        # Send the email with the reset link
        send_reset_email(user.email, uid, token)

        # Send a JSON response
        return Response({"message": "Password reset email sent successfully"}, status=200)

    except User.DoesNotExist:
        # Send a JSON response for error
        return Response({"error": "Email address not found"}, status=404)

@api_view(['POST'])
@permission_classes([])  # Update if permissions are required
def password_reset_confirm(request, uidb64, token):
    """
    Handles password reset confirmation.
    """
    serializer = PasswordResetConfirmSerializer(data=request.data)
    if serializer.is_valid():
        try:
            # Decode the UID and token
            uid = urlsafe_base64_decode(uidb64).decode()
            user = User.objects.get(pk=uid)

            # Validate token
            if not default_token_generator.check_token(user, token):
                return Response({"error": "Invalid or expired token."}, status=400)

            # Set new password
            user.set_password(serializer.validated_data['new_password'])
            user.save()

            return Response({"success": "Password has been reset."})
        except Exception as e:
            return Response({"error": "Invalid token or user not found."}, status=400)
    return Response(serializer.errors, status=400)

# Profile View
class ProfileView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, *args, **kwargs):
        # Retrieve the user's profile
        profile = request.user.profile

        # Serialize the user's profile data (including profile_picture)
        serializer = ProfileSerializer(profile)

        # Return the serialized profile data
        return Response(serializer.data)

    def put(self, request, *args, **kwargs):
        # Get the user's profile
        profile = request.user.profile
        user = request.user
        
        # Update basic user fields
        user.first_name = request.data.get('first_name', user.first_name)
        user.last_name = request.data.get('last_name', user.last_name)
        user.email = request.data.get('email', user.email)

        # Generate OTP secret if not already set
        if not profile.otp_secret:
            import pyotp
            profile.otp_secret = pyotp.random_base32()
            profile.save()

        # Check if a profile picture is being uploaded
        if 'profile_picture' in request.FILES:
            profile.profile_picture = request.FILES['profile_picture']  # Save the uploaded picture
            profile.save()

        user.save()  # Save the user model
        
        return Response({"message": "Profile updated successfully"}, status=status.HTTP_200_OK)
    
# User Registration
class RegisterUserView(APIView):
    permission_classes = [AllowAny]

    def post(self, request):
        serializer = UserRegistrationSerializer(data=request.data)
        if serializer.is_valid():
            # Create the user instance
            user = serializer.save()

            # Create a token for the newly created user
            token, _ = Token.objects.get_or_create(user=user)

            # Return a success response with the token
            return Response({"message": "User registered successfully", "token": token.key}, status=201)
        
        return Response(serializer.errors, status=400)

class RenameFileView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, file_id):
        try:
            # Validate Content-Type
            if request.content_type != 'application/json':
                return Response({"error": "Content-Type must be application/json."}, status=400)

            # Get new name from request
            new_name = request.data.get('name', '').strip()
            if not new_name:
                return Response({"error": "File name cannot be empty."}, status=400)

            # Retrieve file instance
            file_instance = get_object_or_404(File, id=file_id, user_id=request.user.id)
            file_path = os.path.join(settings.MEDIA_ROOT, file_instance.file.name)

            # Check if file exists
            if not os.path.exists(file_path):
                return Response({"error": "File not found."}, status=404)

            # Extract the original file extension
            original_ext = os.path.splitext(file_instance.file.name)[1]

            # Split the new name into base name and extension
            base_name, new_ext = os.path.splitext(new_name)

            # If the user doesn't provide an extension, use the original one
            if not new_ext:
                new_ext = original_ext

            # Generate the final new file name
            new_file_name = f"{slugify(base_name)}{new_ext}"  # Combine slugified base name with extension
            new_file_path = os.path.join(settings.MEDIA_ROOT, 'uploads', new_file_name)

            # Check for duplicate file names
            if os.path.exists(new_file_path):
                return Response({"error": "File with this name already exists."}, status=400)

            # Rename the file
            os.rename(file_path, new_file_path)
            file_instance.file_name = base_name  # Save base name (without extension) in the database
            file_instance.file = f"uploads/{new_file_name}"  # Update file path with extension
            file_instance.save()

            return Response({"message": "File renamed successfully."}, status=200)

        except Exception as e:
            return Response({"error": str(e)}, status=500)

# File Download API
class DownloadFileView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, file_id):
        try:
            # Retrieve the file, ensuring it's the user's file
            file = get_object_or_404(File, id=file_id, user_id=request.user.id)
            encrypted_file_path = file.file.path

            if not os.path.exists(encrypted_file_path):
                return Response({"error": "File not found."}, status=404)

            # Decrypt file to a temporary path
            temp_file_path = os.path.join("/tmp", f"decrypted_{file.file_name}")
            decrypted_path = decrypt_file_to_temp(encrypted_file_path, temp_file_path)

            # Log the download
            AuditLog.objects.create(
                user=request.user,
                action="download",
                file=file,
                timestamp=now()
            )

            # Serve the decrypted file as an attachment
            response = FileResponse(open(decrypted_path, "rb"), as_attachment=True, filename=file.file_name)

            # Delete temporary file after download
            import threading
            def cleanup():
                os.remove(decrypted_path)

            threading.Timer(10, cleanup).start()  # Delete after 10 seconds

            return response

        except File.DoesNotExist:
            return Response({"error": "File not found."}, status=404)
        except Exception as e:
            print(f"Download error: {str(e)}")
            return Response({"error": str(e)}, status=500)

# update new
@require_GET
def get_settings(request):
    # Replace with actual user settings logic
    settings_data = {
        'username': 'example_user',
        'email': 'example@example.com',
        'language': 'en',
    }
    return JsonResponse(settings_data)

@csrf_exempt
def update_username(request):
    if request.method == 'PUT':
        data = json.loads(request.body)
        new_username = data.get('username')
        # Replace with actual logic to update username
        return JsonResponse({'message': f'Username updated to {new_username}'})
    return JsonResponse({'error': 'Invalid request'}, status=400)
    
@api_view(['GET'])
def profile_view(request):
    user = request.user
    profile_data = {
        "username": user.username,
        "email": user.email,
        "first_name": user.first_name,
        "last_name": user.last_name,
        # Add more profile data if needed
    }
    return Response(profile_data)

#show file delete at bin
class DeletedFilesView(APIView):
    permission_classes = [IsAuthenticated]

    # GET method - Fetch deleted files for the user
    def get(self, request):
        deleted_files = DeletedFile.objects.filter(user_id=request.user.id)
        serializer = DeletedFilesSerializer(deleted_files, many=True)
        return Response(serializer.data, status=status.HTTP_200_OK)

    # POST method - Move file to trash
    def post(self, request):
        file_id = request.data.get('file_id')
        if not file_id:
            return Response({"error": "File ID is required."}, status=400)

        try:
            # Fetch the file object
            file = File.objects.get(id=file_id, user_id=request.user.id)

            # Define source and trash paths
            original_path = file.file.path  # Original path
            trash_dir = os.path.join(settings.MEDIA_ROOT, 'uploads', 'trash')  # Trash directory
            os.makedirs(trash_dir, exist_ok=True)  # Ensure trash exists

            # Move file to trash folder
            trash_path = os.path.join(trash_dir, os.path.basename(original_path))
            if os.path.exists(original_path):
                shutil.move(original_path, trash_path)  # Move file physically

            else:
                return Response({"error": "File not found in uploads directory."}, status=404)

            # Save to DeletedFile model
            deleted_file = DeletedFile.objects.create(
                id=file.id,  # Use the same ID
                file=f"uploads/trash/{os.path.basename(file.file.name)}",  # Relative path
                file_name=file.file_name,
                size=file.size,
                user_id=request.user.id,
                deleted_at=timezone.now()
        
            )

            # Remove from active files
            file.delete()

            return Response({"message": "File moved to trash successfully."}, status=200)

        except File.DoesNotExist:
            return Response({"error": "File not found."}, status=404)

        except Exception as e:
            # Rollback: Move the file back if error occurs
            if 'trash_path' in locals() and os.path.exists(trash_path):
                shutil.move(trash_path, original_path)
            return Response({"error": str(e)}, status=500)

# File Delete API
class DeleteFileView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, file_id):
        try:
            # Fetch the file instance
            file_instance = get_object_or_404(File, id=file_id, user_id=request.user.id)

            # Resolve the file path
            file_path = file_instance.file.path
            if not os.path.exists(file_path):
                return Response({"error": "File not found."}, status=404)

            # Move file to trash
            trash_dir = os.path.join(settings.MEDIA_ROOT, 'trash')
            os.makedirs(trash_dir, exist_ok=True)
            trash_path = os.path.join(trash_dir, os.path.basename(file_path))
            os.rename(file_path, trash_path)

            # Add file metadata to DeletedFile table
            DeletedFile.objects.create(
                file=trash_path,
                user_id=request.user.id,
                file_name=file_instance.file_name,
                deleted_at=timezone.now(),
                size=file_instance.size
            )

            # Force delete the file instance (Fallback method)
            File.objects.filter(id=file_instance.id).delete()

            # Confirm file is removed from DB
            if File.objects.filter(id=file_instance.id).exists():
                raise Exception("File record was not deleted properly.")

            return Response({"message": "File moved to trash successfully."}, status=200)

        except Exception as e:
            print(f"Delete error: {str(e)}")
            return Response({"error": str(e)}, status=500)

# File Restore API
class RestoreFileView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        try:
            # Retrieve file IDs from request
            file_ids = request.data.get('file_ids', [])

            if not file_ids:
                return Response({"error": "No file IDs provided."}, status=400)

            restored_files = []
            failed_files = []

            for file_id in file_ids:
                try:
                    # Fetch file metadata
                    deleted_file = get_object_or_404(DeletedFile, id=file_id, user_id=request.user.id)

                    # Move file back from trash
                    trash_path = deleted_file.file
                    restore_dir = os.path.join(settings.MEDIA_ROOT, 'uploads')
                    restored_path = os.path.join(restore_dir, os.path.basename(trash_path))
                    os.rename(trash_path, restored_path)

                    # Restore to File table
                    restored_file = File.objects.create(
                        file=f'uploads/{os.path.basename(restored_path)}',
                        user_id=request.user.id,
                        file_name=deleted_file.file_name,
                        size=deleted_file.size
                    )

                    # Remove from DeletedFile table
                    deleted_file.delete()
                    restored_files.append(restored_file.file_name)

                except Exception as e:
                    failed_files.append({"file_id": file_id, "error": str(e)})

            # Return results
            return Response({
                "restored": restored_files,
                "failed": failed_files
            }, status=200)

        except Exception as e:
            return Response({"error": str(e)}, status=500)

#permenently delete
class PermanentlyDeleteFilesView(APIView):
    permission_classes = [IsAuthenticated]

    def delete(self, request, *args, **kwargs):
        try:
            # Single file deletion by ID
            file_id = kwargs.get('id')
            if file_id:
                # Check if file exists
                deleted_file = get_object_or_404(DeletedFile, id=file_id, user_id=request.user.id)

                # Remove the physical file from the file system
                if os.path.exists(deleted_file.file):
                    os.remove(deleted_file.file)  # Delete file from storage

                # Remove the record from DeletedFile table
                deleted_file.delete()
                return Response({"message": "File permanently deleted."}, status=200)

            # Bulk deletion by IDs from request body
            file_ids = request.data.get('file_ids', [])
            if file_ids:
                deleted_files = DeletedFile.objects.filter(id__in=file_ids, user_id=request.user.id)
                if not deleted_files.exists():
                    return Response({"error": "No valid files found to delete."}, status=404)

                # Delete each file from storage
                for file in deleted_files:
                    if os.path.exists(file.file):
                        os.remove(file.file)

                # Delete records from DeletedFile table
                deleted_files.delete()
                return Response({"message": f"{len(file_ids)} files permanently deleted."}, status=200)

            return Response({"error": "File ID(s) are required."}, status=400)

        except Exception as e:
            return Response({"error": str(e)}, status=500)

#empty bin
class EmptyTrashView(APIView):
    permission_classes = [IsAuthenticated]

    def delete(self, request, *args, **kwargs):
        try:
            # Fetch all deleted files for the user
            deleted_files = DeletedFile.objects.filter(user_id=request.user.id)
            if not deleted_files.exists():
                return Response({"message": "Trash is already empty."}, status=status.HTTP_200_OK)

            # Remove files from the file system
            for deleted_file in deleted_files:
                if os.path.exists(deleted_file.file):
                    os.remove(deleted_file.file)  # Physically delete the file

            # Delete records from database
            deleted_files.delete()

            return Response({"message": "Trash emptied successfully."}, status=status.HTTP_200_OK)

        except Exception as e:
            # Catch any unexpected errors
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

# star file
@api_view(['GET', 'POST'])
@permission_classes([IsAuthenticated])
def starred_files(request):
    if request.method == 'GET':
        # Fetch all starred files for the user
        files = File.objects.filter(user_id=request.user.id, is_starred=True)
        serializer = FileSerializer(files, many=True)
        return Response(serializer.data)

    if request.method == 'POST':
        # Toggle the starred status of a file
        file_id = request.data.get('file_id')
        if not file_id:
            return Response({"error": "File ID is required."}, status=400)
        
        try:
            file = File.objects.get(id=file_id, user=request.user)
            file.is_starred = not file.is_starred  # Toggle the starred status
            file.save()
            return Response({
                "message": f"File '{file.file_name}' starred status updated.",
                "is_starred": file.is_starred
            })
        except File.DoesNotExist:
            return Response({"error": "File not found or not owned by the user."}, status=404)

@csrf_exempt        
def toggle_star(request, id):
    try:
        # Retrieve the file by ID
        file = File.objects.get(id=id)
        
        # Toggle the 'is_starred' status
        file.is_starred = not file.is_starred
        file.save()

        return JsonResponse({'success': True, 'is_starred': file.is_starred})
    except File.DoesNotExist:
        return JsonResponse({'error': 'File not found'}, status=404)

#preview file
class FileView(APIView):
    def get(self, request, file_id):
        try:
            file = get_object_or_404(File, id=file_id)
            encrypted_file_path = file.file.path

            if not os.path.exists(encrypted_file_path):
                return Response({'error': 'File not found'}, status=status.HTTP_404_NOT_FOUND)

            # Decrypt file temporarily
            decrypted_file_path = os.path.join(settings.MEDIA_ROOT, "temp", f"preview_{file.file_name}")
            decrypt_file_to_temp(encrypted_file_path, decrypted_file_path)

            # Determine file type
            file_extension = os.path.splitext(file.file_name)[1].lower()

            if file_extension == '.txt':
                with open(decrypted_file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                os.remove(decrypted_file_path)
                return Response({'name': file.file_name, 'content': content, 'type': 'text'})

            elif file_extension == '.pdf':
                preview_url = request.build_absolute_uri(f"{settings.MEDIA_URL}temp/preview_{file.file_name}")
                return Response({'name': file.file_name, 'type': 'pdf', 'url': preview_url})

            elif file_extension == '.docx':
                document = Document(decrypted_file_path)
                content = "\n\n".join([p.text for p in document.paragraphs])
                tables = [[[cell.text for cell in row.cells] for row in table.rows] for table in document.tables]
                os.remove(decrypted_file_path)
                return Response({'name': file.file_name, 'type': 'docx', 'content': content, 'tables': tables})

            elif file_extension in ['.jpg', '.jpeg', '.png', '.gif']:
                preview_url = request.build_absolute_uri(f"{settings.MEDIA_URL}temp/preview_{file.file_name}")
                return Response({'name': file.file_name, 'type': 'image', 'url': preview_url})

            os.remove(decrypted_file_path)
            return Response({'error': 'Preview not supported for this file type'}, status=400)

        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
def serve_file(request, file_id):
    file = File.objects.get(id=file_id)  # Retrieve the file object from the DB
    file_path = os.path.join(settings.MEDIA_ROOT, file.file_path.name)  # Full path to the file

    if os.path.exists(file_path):  # Check if the file exists on the filesystem
        return FileResponse(open(file_path, 'rb'))  # Serve the file
    else:
        raise Http404("File not found.")  # Return 404 if the file doesn't exist
    
@csrf_exempt
@api_view(['PUT'])
@permission_classes([IsAuthenticated])
def update_profile(request):
    try:
        if request.method == 'PUT':
            # Log incoming request
            print("Incoming request data:", request.body)

            # Parse request body
            data = json.loads(request.body)

            # Access the authenticated user
            user = request.user

            # Update fields (example: username, first_name, last_name)
            if 'username' in data:
                print(f"Updating username to: {data['username']}")
                user.username = data['username']
            if 'first_name' in data:
                print(f"Updating first name to: {data['first_name']}")
                user.first_name = data['first_name']
            if 'last_name' in data:
                print(f"Updating last name to: {data['last_name']}")
                user.last_name = data['last_name']

            # Save the updated user
            user.save()
            print("User saved successfully.")

            # Return success response
            return JsonResponse({
                'message': 'Profile updated successfully',
                'username': user.username,
                'first_name': user.first_name,
                'last_name': user.last_name,
            }, status=200)

    except json.JSONDecodeError:
        print("Error: Invalid JSON format")
        return JsonResponse({'error': 'Invalid JSON format'}, status=400)

    except Exception as e:
        print(f"Unhandled exception: {e}")
        return JsonResponse({'error': str(e)}, status=500)

    # Return 405 if method is not PUT
    return JsonResponse({'error': 'Method not allowed'}, status=405)

@csrf_exempt
@api_view(['POST'])
@permission_classes([IsAuthenticated])
def upload_profile_picture(request):
    if request.method == "POST":
        logger.debug(f"FILES: {request.FILES}")
        if 'profile_picture' in request.FILES:
            file = request.FILES['profile_picture']
            logger.debug(f"Received file: {file.name}")
            # Further file handling logic
            return JsonResponse({"message": "File uploaded successfully"}, status=200)
        else:
            logger.error("No file provided")
            return JsonResponse({"error": "No file provided"}, status=400)
        
class UploadProfilePictureView(APIView):
    parser_classes = (MultiPartParser, FormParser)

    def post(self, request, *args, **kwargs):
        if 'file' not in request.FILES:
            return Response({"error": "No file provided"}, status=status.HTTP_400_BAD_REQUEST)

        profile_picture = request.FILES['file']
        user_profile = request.user.profile
        user_profile.profile_picture = profile_picture
        user_profile.save()

        # Generate the full URL for the uploaded profile picture
        profile_picture_url = f"{settings.MEDIA_URL}{user_profile.profile_picture}"

        return Response(
            {"message": "Profile picture uploaded successfully", "profile_picture_url": profile_picture_url},
            status=status.HTTP_200_OK
        )

# Define pagination first
class FileSearchPagination(PageNumberPagination):
    page_size = 10  # Number of results per page
    page_size_query_param = 'page_size'
    max_page_size = 50

#search function
class FileSearchView(generics.ListAPIView):
    serializer_class = FileSerializer
    filter_backends = [filters.SearchFilter]
    search_fields = ['file_name', 'file_path']
    pagination_class = FileSearchPagination

    def get_queryset(self):
        query = self.request.query_params.get('search', '')
        print(f"Received Search Query: {query}")  # Debugging log
        if query:
            # Filter and order by created_at descending
            return File.objects.filter(
                file_name__icontains=query,
                is_deleted=False
            ).order_by('-created_at')  # Explicit ordering by latest created_at
        return File.objects.none()

#change password
class ChangePasswordView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        user = request.user
        data = request.data

        # Get passwords from request
        old_password = data.get("old_password")
        new_password = data.get("new_password")
        confirm_password = data.get("confirm_password")

        # Check if old password is correct
        if not user.check_password(old_password):
            return Response({"error": "Old password is incorrect."}, status=status.HTTP_400_BAD_REQUEST)

        # Check if new passwords match
        if new_password != confirm_password:
            return Response({"error": "New passwords do not match."}, status=status.HTTP_400_BAD_REQUEST)

        # Validate new password
        try:
            validate_password(new_password, user=user)
        except ValidationError as e:
            return Response({"error": e.messages}, status=status.HTTP_400_BAD_REQUEST)

        # Set new password
        user.set_password(new_password)
        user.save()

        return Response({"message": "Password changed successfully."}, status=status.HTTP_200_OK)

#Delete Account
class DeleteAccountView(APIView):
    permission_classes = [IsAuthenticated]

    def delete(self, request):
        user = request.user

        # Remove associated profile picture if exists
        profile = user.profile  # Assuming User has a related Profile model
        if profile.profile_picture:
            profile_picture_path = profile.profile_picture.path
            if os.path.exists(profile_picture_path):
                os.remove(profile_picture_path)

        # Delete the user and related data
        user.delete()

        return Response({"message": "Account deleted successfully"}, status=status.HTTP_200_OK)

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def verify_2fa(request):
    try:
        data = request.data
        otp = data.get('otp')
        
        if not otp:
            logger.error("No OTP provided in request")
            return JsonResponse({'error': 'OTP is required'}, status=400)

        # Log the received OTP
        logger.info(f"Received OTP: {otp}")

        user = request.user
        base32_secret = user.profile.otp_secret
        
        if not base32_secret:
            logger.error("No OTP secret available for the user")
            return JsonResponse({'error': 'No OTP secret configured'}, status=400)

        # Log the secret being used
        logger.info(f"Using secret: {base32_secret}")

        totp = pyotp.TOTP(base32_secret)

        # Generate the current OTP on the server and log it
        server_otp = totp.now()
        logger.info(f"Generated OTP on Server: {server_otp}")

        # Validate the received OTP
        if totp.verify(otp, valid_window=1):  # ±30 seconds drift allowed
            logger.info("OTP is valid")
            return JsonResponse({'success': True, 'message': '2FA verified successfully'})
        else:
            logger.error("Invalid OTP provided")
            return JsonResponse({'error': 'Invalid OTP'}, status=400)

    except Exception as e:
        logger.error("Error during OTP verification", exc_info=True)
        return JsonResponse({'error': 'Server error'}, status=500)

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def setup_2fa(request):
    user = request.user
    profile = user.profile

    if not profile.otp_secret:
        profile.otp_secret = pyotp.random_base32()  # Generate Base32 secret
        profile.save()

    base32_secret = profile.otp_secret

    otp_url = pyotp.totp.TOTP(base32_secret).provisioning_uri(
        name=user.username,
        issuer_name="MyApp"  # Replace with your app's name
    )

    qr = qrcode.QRCode(version=1, error_correction=qrcode.constants.ERROR_CORRECT_L, box_size=10, border=4)
    qr.add_data(otp_url)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")

    buffer = BytesIO()
    img.save(buffer, format="PNG")
    buffer.seek(0)
    img_str = base64.b64encode(buffer.getvalue()).decode('utf-8')

    return JsonResponse({
        "qr_code": f"data:image/png;base64,{img_str}",
        "otp_url": otp_url,
        "secret": base32_secret  # Optional for manual setup
    })

class Enable2FAView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        # Get the user's profile
        profile = request.user.profile

        # Check if OTP secret already exists
        if not profile.otp_secret:
            # Generate and save OTP secret
            profile.otp_secret = pyotp.random_base32()
            profile.save()

        # Return the secret to the user to display QR code
        return Response({"otp_secret": profile.otp_secret}, status=status.HTTP_200_OK)
    
@receiver(post_save, sender=User)
def create_or_update_user_profile(sender, instance, created, **kwargs):
    if created:
        Profile.objects.create(user=instance)
    instance.profile.save()

#testing email
def send_test_email(request):
    """Django view to send a test email."""
    try:
        to_email = 'recipient@example.com'
        subject = 'Test Email from Django'
        body = 'This is a test email sent via the Gmail API.'

        # Send the email
        result = send_mail(to_email, subject, body)

        return JsonResponse({'status': 'success', 'message': result})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)})

#share the file
class ShareFileView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, file_id):
        logger.info("Received request to share file")
        email = request.data.get('email')
        password = request.data.get('password')  # Get optional password
        allow_download = request.data.get('allow_download', False ) 
        one_time_view = request.data.get('one_time_view', False) 

        logger.info(f"Sharing file with one_time_view: {one_time_view}")

        if not email:
            return Response({"error": "Email is required."}, status=400)

        # Validate email
        try:
            validate_email(email)
        except ValidationError:
            return Response({"error": "Invalid email address."}, status=400)

        # Check if file exists and belongs to the current user
        file_instance = get_object_or_404(File, id=file_id, user_id=request.user.id)

        # Generate a secure, expiring share link
        share_token = str(uuid.uuid4())
        expiry_time = now() + timedelta(hours=24)  # Link expires in 24 hours
        
        shared_file = SharedFile.objects.create(
            file=file_instance,
            shared_with=email,
            shared_by=request.user,
            share_token=share_token,
            expiry_time=expiry_time,
            allow_download=allow_download,
            one_time_view=one_time_view
        )

         # 🔹 Set password if provided
        if password:
            shared_file.set_password(password)
        shared_file.save()

        # Notify the recipient
        try:
            subject = "File Shared with You"
            body = (
                f"{request.user.username} has shared a file with you: {file_instance.file_name}. "
                f"This link will expire in 24 hours.\n\n"
                f"Access the file here: https://127.0.0.1:8000/api/shared-files/{share_token}/"
            )
            #if password:
             #   body += f" Password: {password}\n"

            send_email_via_gmail(email, subject, body)
        except Exception as e:
            logger.error(f"Error sending email: {e}")
            return Response({"error": "Failed to send email notification."}, status=500)

        return Response({
            "message": "File shared successfully!",
            "share_link": f"https://127.0.0.1:8000/api/shared-files/{share_token}/"
        }, status=200)

#expired link    
class RetrieveSharedFileView(APIView):
    def get(self, request, share_token):
        shared_file = get_object_or_404(SharedFile, share_token=share_token)

        if shared_file.is_expired():
            return Response({"error": "This link has expired."}, status=400)
        
        # Get the entered password from the request
        entered_password = request.GET.get('password', '')

        if shared_file.password_hash and not shared_file.check_password(entered_password):
            return Response({"error": "Incorrect password."}, status=403)

        # 🔹 If the file was shared with a password, validate it
        if shared_file.password_hash:
            if not entered_password:
                return Response({"error": "Password is required."}, status=401)

        # 🔹 Check if it's a one-time view and has been accessed before
        if shared_file.one_time_view and shared_file.has_been_viewed:
            return Response({"error": "This file has already been viewed and cannot be accessed again."}, status=403)

        # 🔹 If it's a one-time view, mark it as viewed
        if shared_file.one_time_view:
            shared_file.has_been_viewed = True
            shared_file.save(update_fields=['has_been_viewed']) 

        #Check if user can download the file
        can_download = shared_file.can_download

        return Response({
            "message": "File link is valid.",
            "file_name": shared_file.file.file_name,
            "file_url": request.build_absolute_uri(shared_file.file.file.url),
            "allow_download": can_download
        })

#view file share from link
class SharedFileView(View):
    def get(self, request, share_token):
        temp_file_path = None

        try:
            shared_file = SharedFile.objects.get(share_token=share_token)
            
            # Check if the link has expired
            if shared_file.is_expired():
                return HttpResponse("This link has expired.", status=410)  
            
            # 🔹 Enforce one-time view
            if shared_file.one_time_view and shared_file.has_been_viewed:
                return HttpResponse("This file has already been viewed and cannot be accessed again.", status=403)
            
            # Handle password protection
            password = request.GET.get("password", None)
            if shared_file.password_hash and not shared_file.check_password(password):
                return render(request, "enter_password.html", {
                    "share_token": share_token,
                    "error": "Incorrect password. Please try again."
                })
            
            # Get the file path
            file_path = shared_file.file.file.path  # Full file path
            file_name = shared_file.file.file_name  # File name from the database
            
            # Check if the file exists
            if not os.path.exists(file_path):
                return HttpResponse("File not found.", status=404)
            
            # 🔹 Decrypt file to a temporary location
            temp_file_path = f"/tmp/decrypted_{file_name}"
            decrypt_file_to_temp(file_path, temp_file_path)
            
            # 🔹 Determine MIME type
            mime_type, _ = mimetypes.guess_type(file_name)
            if not mime_type:
                mime_type = "application/octet-stream"  # Default type if unknown
            
             # 🔹 Serve the decrypted file
            with open(temp_file_path, 'rb') as temp_file:
                response = HttpResponse(temp_file.read(), content_type=mime_type)
                response['Content-Disposition'] = f'inline; filename="{file_name}"'

                # 🔹 Mark file as viewed if one-time view is enabled
                if shared_file.one_time_view:
                    shared_file.has_been_viewed = True
                    shared_file.save(update_fields=['has_been_viewed'])  # ✅ Save

                return response
            
        except SharedFile.DoesNotExist:
            return HttpResponse("Invalid or expired link.", status=400)

        except Exception as e:
            return HttpResponse(f"Error serving file: {str(e)}", status=500)

        finally:
            if temp_file_path and os.path.exists(temp_file_path):  # ✅ Ensure file exists
                os.remove(temp_file_path)
            
def extract_text_from_docx(docx_file_path):
    doc = Document(docx_file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def move_to_lockbox(request, file_id):
    try:
        file = File.objects.get(id=file_id, user=request.user)
        if file.is_locked:
            return Response({'error': 'File is already in Lock Box!'}, status=400)
        
        file.is_locked = True  # Move to Lock Box
        file.save()
        
        return Response({'message': 'File moved to Lock Box successfully!'})
    except File.DoesNotExist:
        return Response({'error': 'File not found!'}, status=404)

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def move_out_of_lockbox(request, file_id):
    try:
        file = File.objects.get(id=file_id, user=request.user)
        file.is_locked = False
        file.save()
        return Response({'message': 'File moved out of Lock Box successfully!'})
    except File.DoesNotExist:
        return Response({'error': 'File not found!'}, status=404)

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def get_locked_files(request):
    if not request.user.is_authenticated:
        return JsonResponse({'error': 'Unauthorized access! Please log in.'}, status=401)

    files = File.objects.filter(user=request.user, is_locked=True)
    return Response([{'id': f.id, 'name': f.file_name, 'created_at': f.created_at} for f in files])

@csrf_exempt  # ✅ To avoid CSRF issues
@api_view(['POST'])
@permission_classes([IsAuthenticated]) 
def verify_lockbox_password(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        enteredPassword = data.get('password')
        user = request.user

        try:
            lockbox = Lockbox.objects.get(user=user)
            if lockbox.check_password(enteredPassword):  # ✅ Make sure `check_password` is implemented
                return JsonResponse({'success': True})
            else:
                return JsonResponse({'success': False, 'error': 'Incorrect password'}, status=400)
        except Lockbox.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Lockbox not found'}, status=404)
    
@csrf_exempt
@api_view(['POST'])
@permission_classes([IsAuthenticated])  
def save_lockbox_password(request):
    data = json.loads(request.body)
    new_password = data.get('password')
    user = request.user

    if not new_password:
        return JsonResponse({'success': False, 'error': 'Password is required'}, status=400)

    lockbox, created = Lockbox.objects.get_or_create(user=user)
    lockbox.set_password(new_password)  # ✅ Hash & save password

    return JsonResponse({'success': True, 'message': 'Password saved successfully'})